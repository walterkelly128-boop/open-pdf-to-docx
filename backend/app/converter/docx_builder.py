from io import BytesIO
from pathlib import Path
from typing import Iterable
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .layout import LayoutBlock, LayoutLine, build_layout
from .pdf_parser import PdfImage, PdfPage, PdfWord, iter_pages

_FONT_MAP = {
    "ArialMT": "Arial",
    "Arial-BoldMT": "Arial",
    "TimesNewRomanPSMT": "Times New Roman",
    "TimesNewRomanPS-BoldMT": "Times New Roman",
    "Helvetica": "Arial",
    "Helvetica-Bold": "Arial",
    "Helvetica-Oblique": "Arial",
    "Courier": "Courier New",
    "Calibri": "Calibri",
    "SimSun": "SimSun",
    "宋体": "SimSun",
    "Microsoft YaHei": "Microsoft YaHei",
    "微软雅黑": "Microsoft YaHei",
}

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
V_NS = "urn:schemas-microsoft-com:vml"
O_NS = "urn:schemas-microsoft-com:office:office"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"


def safe_font(name: str) -> str:
    base = name.split("+")[-1].strip()
    if base in _FONT_MAP:
        return _FONT_MAP[base]
    for suffix in ("-BoldItalic", "-Bold", "-Italic", "_Bold", "_Italic"):
        if base.endswith(suffix):
            base = base[: -len(suffix)]
            break
    return _FONT_MAP.get(base, base or "Arial")


def _set_run_font(run, word: PdfWord) -> None:
    name = safe_font(word.font)
    run.font.name = name
    run.font.size = Pt(max(1.0, word.size))
    run.bold = word.bold
    run.italic = word.italic
    value = max(0, min(0xFFFFFF, int(word.color)))
    run.font.color.rgb = RGBColor((value >> 16) & 255, (value >> 8) & 255, value & 255)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), name)


def _set_paragraph_exact_line_height(paragraph, height_pt: float) -> None:
    spacing = paragraph.paragraph_format
    spacing.space_before = Pt(0)
    spacing.space_after = Pt(0)
    ppr = paragraph._p.get_or_add_pPr()
    el = ppr.find(qn("w:spacing"))
    if el is None:
        el = OxmlElement("w:spacing")
        ppr.append(el)
    el.set(qn("w:line"), str(max(1, round(height_pt * 20))))
    el.set(qn("w:lineRule"), "exact")


def _add_text_box(doc: Document, block: LayoutBlock, page: PdfPage, shape_id: int) -> None:
    width = max(4.0, block.x1 - block.x0 + 2.0)
    height = max(4.0, block.y1 - block.y0 + 3.0)
    style = (
        f"position:absolute;margin-left:{block.x0:.2f}pt;margin-top:{block.y0:.2f}pt;"
        f"width:{width:.2f}pt;height:{height:.2f}pt;"
        "mso-position-horizontal-relative:page;"
        "mso-position-vertical-relative:page;"
        f"z-index:{100 + shape_id};mso-wrap-style:none;"
    )
    xml = (
        f'<w:pict xmlns:w="{W_NS}" xmlns:v="{V_NS}" xmlns:o="{O_NS}" xmlns:r="{R_NS}">'
        f'<v:shape id="TextBox{shape_id}" type="#_x0000_t202" '
        f'style="{escape(style)}" stroked="f" filled="f">'
        '<v:textbox inset="0,0,0,0"><w:txbxContent>'
        '</w:txbxContent></v:textbox></v:shape></w:pict>'
    )
    pict = parse_xml(xml)
    txbx = pict.xpath('.//w:txbxContent', namespaces={"w": W_NS})[0]

    for line in block.lines:
        p = OxmlElement("w:p")
        ppr = OxmlElement("w:pPr")
        p.append(ppr)
        spacing = OxmlElement("w:spacing")
        line_height = max(line.height, max((w.size for w in line.words), default=8) * 1.05)
        spacing.set(qn("w:line"), str(max(1, round(line_height * 20))))
        spacing.set(qn("w:lineRule"), "exact")
        ppr.append(spacing)
        previous = None
        for word in line.words:
            r = OxmlElement("w:r")
            rpr = OxmlElement("w:rPr")
            r.append(rpr)
            rfonts = OxmlElement("w:rFonts")
            name = safe_font(word.font)
            for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
                rfonts.set(qn(f"w:{attr}"), name)
            rpr.append(rfonts)
            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), str(max(2, round(word.size * 2))))
            rpr.append(sz)
            if word.bold:
                rpr.append(OxmlElement("w:b"))
            if word.italic:
                rpr.append(OxmlElement("w:i"))
            color = OxmlElement("w:color")
            value = max(0, min(0xFFFFFF, int(word.color)))
            color.set(qn("w:val"), f"{value:06X}")
            rpr.append(color)
            text = word.text
            if previous is not None:
                gap = word.x0 - previous.x1
                threshold = max(1.5, min(previous.size, word.size) * 0.18)
                if gap > threshold and previous.text[-1:].isascii() and word.text[:1].isascii():
                    text = " " + text
            t = OxmlElement("w:t")
            if text[:1].isspace() or text[-1:].isspace():
                t.set(qn("xml:space"), "preserve")
            t.text = text
            r.append(t)
            p.append(r)
            previous = word
        txbx.append(p)

    doc.paragraphs[-1]._p.append(pict)


def _add_positioned_image(doc: Document, image: PdfImage, page: PdfPage, shape_id: int) -> None:
    try:
        r_id, _ = doc.part.get_or_add_image(BytesIO(image.data))
    except Exception:
        return
    width = max(2.0, image.x1 - image.x0)
    height = max(2.0, image.y1 - image.y0)
    style = (
        f"position:absolute;margin-left:{image.x0:.2f}pt;margin-top:{image.y0:.2f}pt;"
        f"width:{width:.2f}pt;height:{height:.2f}pt;"
        "mso-position-horizontal-relative:page;"
        "mso-position-vertical-relative:page;"
        f"z-index:{shape_id};mso-wrap-style:none;"
    )
    xml = (
        f'<w:pict xmlns:w="{W_NS}" xmlns:v="{V_NS}" xmlns:o="{O_NS}" xmlns:r="{R_NS}">'
        f'<v:shape id="Image{shape_id}" type="#_x0000_t75" style="{escape(style)}" '
        'stroked="f" filled="f"><v:imagedata r:id="' + r_id + '" o:title="PDF image" />'
        '</v:shape></w:pict>'
    )
    doc.paragraphs[-1]._p.append(parse_xml(xml))


def _configure_section(section, page: PdfPage) -> None:
    section.page_width = Inches(page.width / 72)
    section.page_height = Inches(page.height / 72)
    section.top_margin = section.bottom_margin = Inches(0)
    section.left_margin = section.right_margin = Inches(0)
    section.header_distance = Inches(0)
    section.footer_distance = Inches(0)


def _content_items(page: PdfPage) -> Iterable[tuple[float, int, object]]:
    blocks = build_layout(page)
    items: list[tuple[float, int, object]] = [(b.y0, 1, b) for b in blocks]
    items.extend((i.y0, 0, i) for i in page.images)
    return sorted(items, key=lambda item: (item[0], item[1]))


def _add_page_anchor_paragraph(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    return p


def _convert_positioned_editable(pdf_path: str | Path, output_path: str | Path) -> None:
    """Rebuild each PDF page as independently editable, absolutely positioned Word objects.

    Text remains real DOCX text inside editable Word text boxes. Images are real
    embedded Word images. Coordinates are preserved relative to the PDF page,
    avoiding the large drift caused by ordinary flowing paragraphs.
    """
    doc = Document()
    first = True
    shape_id = 1
    for page in iter_pages(pdf_path):
        section = doc.sections[0] if first else doc.add_section(WD_SECTION.NEW_PAGE)
        first = False
        _configure_section(section, page)
        anchor = _add_page_anchor_paragraph(doc)
        for _, kind, item in _content_items(page):
            if kind == 0:
                _add_positioned_image(doc, item, page, shape_id)
            else:
                _add_text_box(doc, item, page, shape_id)
            shape_id += 1
        # Keep a minimal page anchor so the section has a stable body paragraph.
        anchor.paragraph_format.space_after = Pt(0)

    if len(doc.paragraphs) > 1 and not doc.paragraphs[-1].text and len(doc.paragraphs[-1]._p) == 1:
        element = doc.paragraphs[-1]._element
        element.getparent().remove(element)
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def _convert_editable(pdf_path: str | Path, output_path: str | Path) -> None:
    """Compatibility editable mode; uses the same positioned text reconstruction."""
    _convert_positioned_editable(pdf_path, output_path)


def convert_pdf_to_docx(
    pdf_path: str | Path,
    output_path: str | Path,
    mode: str = "fidelity",
    fidelity_dpi: int = 180,
) -> None:
    mode = (mode or "fidelity").lower().strip()
    if mode not in {"fidelity", "editable"}:
        raise ValueError("mode must be 'fidelity' or 'editable'")
    # Both modes now preserve real editable text. fidelity means coordinate-first
    # reconstruction rather than rasterization; the DPI argument is retained for API compatibility.
    _convert_positioned_editable(pdf_path, output_path)
